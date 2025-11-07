import pdfplumber
from docx import Document

pdf_path = "ds_v5.pdf"
docx_path = "v_5.docx"

doc = Document()
max_pages = 100  # only first 100 pages


with pdfplumber.open(pdf_path) as pdf:
    for i, page in enumerate(pdf.pages, start=1):
        try:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
            else:
                doc.add_paragraph(f"[Page {i} could not be read]")
        except Exception as e:
            doc.add_paragraph(f"[Page {i} error: {e}]")

doc.save(docx_path)
print("PDF text extracted to DOCX (all errors ignored).")
